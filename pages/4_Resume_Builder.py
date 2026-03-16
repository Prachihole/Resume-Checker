## ====================== IMPORTS ======================
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from io import BytesIO
from PIL import Image
import language_tool_python
import os

ROLE_KEYWORDS = {
    "AI / ML Engineer": [
        "python","machine learning","deep learning","computer vision",
        "nlp","tensorflow","pytorch","model training","deployment"
    ],
    "Data Scientist": [
        "python","statistics","data analysis","machine learning",
        "pandas","numpy","scikit-learn","visualization"
    ],
    "Web Developer": [
        "html","css","javascript","react","node","api","frontend","backend"
    ],
    "Software Engineer": [
        "python","java","c++","data structures","algorithms",
        "problem solving","system design"
    ]
}

# ====================== PAGE CONFIG (FIRST CALL) ======================
st.set_page_config(
    page_title=" Smart Resume Builder",
    page_icon="🧠",
    layout="wide"
)

# ====================== LOGIN CHECK ======================
if not st.session_state.get("logged_in", False):
    st.warning("⚠️ Please log in to access this page.")
    st.stop()

# ====================== CACHED GRAMMAR TOOL ======================
@st.cache_resource
def load_tool():
    return language_tool_python.LanguageToolPublicAPI("en-US")
grammar_tool = load_tool()

# ====================== HEADER ======================
st.markdown("""
<h1 style='text-align:center;
background: linear-gradient(90deg,#00FFFF,#FF63E0);
-webkit-background-clip:text;
-webkit-text-fill-color:transparent;'>
🧠 Smart Resume Builder 
</h1>
""", unsafe_allow_html=True)

st.caption("Build a clean, one-page, ATS-friendly resume from scratch 🚀")

# ====================== BASIC INFO ======================
with st.container(border=True):
    col1, col2 = st.columns([2,1])

    with col1:
        name = st.text_input("👤 Full Name", st.session_state.get("user_name",""))
        job_role = st.selectbox(
            "💼 Target Role",
            list(ROLE_KEYWORDS.keys())
)

        email = st.text_input("✉️ Email", st.session_state.get("user_email",""))
        phone = st.text_input("📞 Phone Number")
        linkedin = st.text_input("🔗 LinkedIn URL")
        github = st.text_input("💻 GitHub URL")

    with col2:
        st.markdown("📸 **Profile Photo (optional)**")
        photo = st.file_uploader("Upload JPG / PNG", type=["jpg","jpeg","png"])

        if photo and photo.size > 2 * 1024 * 1024:
           st.error("Image too large. Upload under 2MB.")
           st.stop()
        img_path = None
        if photo:
            image = Image.open(photo)
            image.thumbnail((150,150))
            st.image(image)
            img_path = f"profile_{email}.png"
            img_bytes = BytesIO()
            image.save(img_bytes, format="PNG")
            img_bytes.seek(0)

# ====================== SECTION TOGGLES ======================
st.markdown("## 🧩 Resume Sections")

sections_enabled = {
    "Professional Summary": st.checkbox("Professional Summary", True),
    "Skills": st.checkbox("Skills", True),
    "Experience": st.checkbox("Experience", True),
    "Projects": st.checkbox("Projects", True),
    "Education": st.checkbox("Education", True)
}

# ====================== MAIN SECTIONS ======================
summary = st.text_area("🧠 Professional Summary") if sections_enabled["Professional Summary"] else ""
st.caption("💡 2–3 lines. Avoid 'I am'. Focus on impact.")

skills = st.text_area("🧩 Skills (comma separated)") if sections_enabled["Skills"] else ""
st.caption("💡 ATS prefers keywords, not sentences.")

experience = st.text_area("🏢 Experience") if sections_enabled["Experience"] else ""
st.caption("💡 Use action verbs + results.")

projects = st.text_area("💼 Projects") if sections_enabled["Projects"] else ""
education = st.text_area("🎓 Education") if sections_enabled["Education"] else ""

# ====================== CUSTOM SECTIONS ======================
st.markdown("## ➕ Custom Sections")

if "custom_sections" not in st.session_state:
    st.session_state.custom_sections = []

new_section = st.text_input(
    "Add a new section (e.g. Certifications, Achievements, Publications)"
)

if st.button("➕ Add Section"):
    if new_section and new_section not in st.session_state.custom_sections:
        st.session_state.custom_sections.append(new_section)

for sec in st.session_state.custom_sections:
    st.text_area(f"✏️ {sec}", key=f"custom_{sec}")

# ====================== THEME ======================
theme = st.selectbox("🎨 Resume Theme", ["Elegant Blue", "Minimal Gray"])
heading_color = RGBColor(0,102,255) if "Blue" in theme else RGBColor(90,90,90)


def clean_lines(text):
    lines = []
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        # remove leading bullets or dashes
        line = line.lstrip("•-– ").strip()
        lines.append(line)
    return lines


# ====================== COMPLETENESS SCORE ======================
sections_text = {
    "Summary": summary,
    "Skills": skills,
    "Experience": experience,
    "Projects": projects,
    "Education": education,
    **{sec: st.session_state.get(f"custom_{sec}", "") for sec in st.session_state.custom_sections}
}

filled = sum(1 for v in sections_text.values() if v and len(v.strip()) > 30)
completion = int((filled / max(len(sections_text),1)) * 100)

st.markdown("## 📈 Resume Completeness")
st.progress(completion / 100)
st.caption(f"{completion}% complete — recruiters prefer 85%+")

# ====================== DOCX GENERATION ======================
def add_section(doc, title, content, mode="bullet"):
    if not content.strip():
        return

    heading = doc.add_heading(title.upper(), level=1)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)

    lines = clean_lines(content)

    if mode == "skills":
        # compact skills layout
        doc.add_paragraph(", ".join(lines))
        return

    if mode == "projects":
        for line in lines:
            if "—" in line or ":" in line:
                p = doc.add_paragraph(line)
                p.runs[0].bold = True
            else:
                doc.add_paragraph(line, style="List Bullet")
        return

    # default bullet mode
    for line in lines:
        doc.add_paragraph(line, style="List Bullet")


def create_resume():
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Inches(0.5)
        s.left_margin = s.right_margin = Inches(0.6)

    # Header
   # ====================== HEADER WITH PHOTO (PROFESSIONAL LAYOUT) ======================
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

# Column widths
    left_cell.width = Inches(5.5)
    right_cell.width = Inches(1.5)

# LEFT SIDE: Name, title, contact
    p = left_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    name_run = p.add_run(name + "\n")
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = heading_color

    title_run = p.add_run(job_role + "\n")

    title_run.font.size = Pt(12)
    contact_run = p.add_run(f"{email} | {phone}\n{linkedin} | {github}")
    contact_run.font.size = Pt(10)

# RIGHT SIDE: Photo (top-right)
    if img_path and os.path.exists(img_path):
       pr = right_cell.paragraphs[0]
       pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
       pr.add_run().add_picture(img_path, width=Inches(1.0))


    add_section(doc, "Professional Summary", summary)
    add_section(doc, "Skills", skills, mode="skills")
    add_section(doc, "Experience", experience)
    add_section(doc, "Projects", projects, mode="projects")
    add_section(doc, "Education", education)

    for sec in st.session_state.custom_sections:
       add_section(doc, sec, st.session_state.get(f"custom_{sec}", ""))

    

    doc.add_paragraph(
        f"\nGenerated on {datetime.now().strftime('%d %b %Y')} • Smart Resume Builder 4.0"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ====================== ATS + GRAMMAR ======================
def ats_analysis(text, role):
    base_score = 50
    keywords = ROLE_KEYWORDS.get(role, [])
    found = [k for k in keywords if k in text.lower()]

    score = base_score + len(found) * 4

    if len(text.split()) > 200:
        score += 5

    return min(score, 100), found



def clean_grammar_messages(issues, max_items=4):
    seen = set()
    cleaned = []

    for issue in issues:
        msg = issue.message.strip()

        # normalize spammy messages
        if "Possible spelling mistake" in msg:
            msg = "Possible spelling mistakes detected"
        if "whitespace" in msg.lower():
            msg = "Extra or repeated spaces found"

        if msg not in seen:
            seen.add(msg)
            cleaned.append(msg)

        if len(cleaned) >= max_items:
            break

    return cleaned

# ====================== GENERATE ======================
st.markdown("---")

if st.button("🚀 Generate Resume & Analyze", use_container_width=True):

    if not name or not email or not phone:
        st.warning("Please fill all basic details.")
    else:
        combined_text = " ".join(sections_text.values())
        total_words = len(combined_text.split())

        ats_score, found_keys = ats_analysis(combined_text, job_role)

        @st.cache_data
        def check_grammar(text):
            return grammar_tool.check(text)
        grammar_issues = check_grammar(combined_text)[:8]
        raw_issues = check_grammar(combined_text)
        st.success(f"✅ Resume Generated — ATS Score: {ats_score}%")

        st.markdown("## 🧠 Resume Intelligence")
        c1, c2, c3 = st.columns(3)
        c1.metric("ATS Score", f"{ats_score}%")
        c2.metric("Keywords Matched", len(found_keys))
        c3.metric("Grammar Issues", len(grammar_issues))

        st.markdown("## 📄 One-Page Check")
        st.write(f"Word count: **{total_words}**")

        if total_words < 250:
            st.info(
        "ℹ️ Resume is short. Consider adding:\n"
        "- 2–3 impact points in Experience\n"
        "- Tools & technologies in Skills\n"
        "- Metrics (accuracy, %, performance)"
    )
        elif total_words > 700:
           st.warning("⚠️ Resume may exceed one page. Consider trimming.")
        else:
           st.success("✅ Ideal one-page resume length")



        grammar_hints = clean_grammar_messages(raw_issues)

        st.markdown("## ✍️ Grammar Hints")

        for g in grammar_hints:
            st.markdown(f"- {g}")

        if len(raw_issues) > len(grammar_hints):
           st.caption(f"➕ {len(raw_issues) - len(grammar_hints)} more minor issues hidden")

        st.markdown("### 👀 Resume Preview (Text)")
        st.text_area(
              "Preview",
               combined_text,
               height=250,
               disabled=True
)

        st.download_button(
            "⬇️ Download Resume (DOCX)",
            data=create_resume(),
            file_name=f"{name}_Resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
